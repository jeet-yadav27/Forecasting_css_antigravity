"""
Generate multivariate model architecture diagrams and a Word (.docx) export.

Run from project root:
    python scripts/generate_model_architecture_docs.py
"""
from __future__ import annotations

import os
import sys

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIAG_DIR = os.path.join(PROJECT_ROOT, "docs", "diagrams")
DOCX_PATH = os.path.join(PROJECT_ROOT, "docs", "MULTIVARIATE_MODEL_ARCHITECTURE.docx")

sys.path.insert(0, PROJECT_ROOT)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def _box(ax, x, y, w, h, text, color, fontsize=8):
    patch = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.02,rounding_size=0.08",
        facecolor=color, edgecolor="#1a1a2e", linewidth=1.1,
    )
    ax.add_patch(patch)
    ax.text(
        x + w / 2, y + h / 2, text, ha="center", va="center",
        fontsize=fontsize, color="#111", wrap=True, fontweight="normal",
    )


def _stack_diagram(path: str, title: str, layers: list[tuple[str, str]], width=8.2, height=None):
    n = len(layers)
    height = height or max(5.2, 0.72 * n + 1.4)
    fig, ax = plt.subplots(figsize=(width, height))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, n + 1.2)
    ax.axis("off")
    ax.set_title(title, fontsize=13, fontweight="bold", pad=8, color="#1a1a2e")
    for i, (label, color) in enumerate(reversed(layers)):
        y = i + 0.35
        _box(ax, 1.2, y, 7.6, 0.78, label, color, fontsize=8.2)
        if i < n - 1:
            ax.annotate(
                "", xy=(5, y + 0.82), xytext=(5, y + 0.95),
                arrowprops=dict(arrowstyle="->", color="#444", lw=1.2),
            )
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def generate_diagrams() -> dict[str, str]:
    os.makedirs(DIAG_DIR, exist_ok=True)
    paths = {}

    # Pipeline
    fig, ax = plt.subplots(figsize=(11.2, 6.4))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title(
        "Multivariate forecasting pipeline (Tab 1)",
        fontsize=13, fontweight="bold", color="#1a1a2e",
    )
    blocks = [
        (0.3, 5.6, 2.6, 1.6, "Upload claims\n+ production", "#E3F2FD"),
        (3.2, 5.6, 2.8, 1.6, "loader.py\nmonthly series\n+ EXOG fusion", "#C8E6C9"),
        (6.3, 5.6, 2.8, 1.6, "MinMaxScaler\nclaims + exog\n(leak-free / fold)", "#FFF9C4"),
        (9.4, 5.6, 4.2, 1.6, "Windows X:(N,W,1+F)\nY:(N,horizon)", "#FFE0B2"),
        (0.3, 3.0, 3.0, 1.8, "CNN-LSTM\nConv→LSTM→Dense", "#D1C4E9"),
        (3.5, 3.0, 3.0, 1.8, "N-BEATS\n6 residual blocks", "#F8BBD0"),
        (6.7, 3.0, 3.0, 1.8, "Transformer\nMHA + FFN residual", "#B2DFDB"),
        (9.9, 3.0, 3.7, 1.8, "SARIMA\nunivariate claims\n(scaled)", "#DCEDC8"),
        (1.5, 0.4, 4.4, 1.8, "Rank RMSE/MAE/MAPE/R²\nselect best model", "#BBDEFB"),
        (6.4, 0.4, 6.6, 1.8, "Inverse-MAE ensemble + CI\nCM multiply → UI / pickle", "#FFCCBC"),
    ]
    for x, y, w, h, t, c in blocks:
        _box(ax, x, y, w, h, t, c, fontsize=8)
    ax.annotate("", xy=(3.2, 6.4), xytext=(2.9, 6.4),
                arrowprops=dict(arrowstyle="->", lw=1.4))
    ax.annotate("", xy=(6.3, 6.4), xytext=(6.0, 6.4),
                arrowprops=dict(arrowstyle="->", lw=1.4))
    ax.annotate("", xy=(9.4, 6.4), xytext=(9.1, 6.4),
                arrowprops=dict(arrowstyle="->", lw=1.4))
    ax.annotate("", xy=(7.0, 4.8), xytext=(7.7, 5.55),
                arrowprops=dict(arrowstyle="->", lw=1.2))
    ax.annotate("", xy=(3.7, 2.2), xytext=(5.5, 3.0),
                arrowprops=dict(arrowstyle="->", lw=1.2))
    ax.annotate("", xy=(9.5, 2.2), xytext=(8.5, 3.0),
                arrowprops=dict(arrowstyle="->", lw=1.2))
    fig.tight_layout()
    p = os.path.join(DIAG_DIR, "pipeline_multivariate.png")
    fig.savefig(p, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    paths["pipeline"] = p

    _stack_diagram(
        os.path.join(DIAG_DIR, "cnn_lstm.png"),
        "CNN-LSTM — layer stack (NumPy, Tab 1)",
        [
            ("Input  (W, F)  W=lookback, F=1+len(EXOG)  fused claims + exog", "#E3F2FD"),
            ("Preprocess  MinMax claims & exog  (pipeline, not inside net)", "#BBDEFB"),
            ("Conv1D-1  k=3,  F→16  valid pad  + ReLU  + inverted dropout", "#C5CAE9"),
            ("Conv1D-2  k=3, 16→8  valid pad  + ReLU  + inverted dropout", "#9FA8DA"),
            ("LSTM  hidden=32  4 gates (i,f,g,o)  forget bias = 1", "#7986CB"),
            ("Dropout on final hidden h  (train only)", "#5C6BC0"),
            ("Dense  32 → horizon  linear  ŷ", "#3F51B5"),
        ],
    )
    paths["cnn"] = os.path.join(DIAG_DIR, "cnn_lstm.png")

    _stack_diagram(
        os.path.join(DIAG_DIR, "nbeats.png"),
        "N-BEATS — block + stack (NumPy, Tab 1)",
        [
            ("Input flatten  (W, F) → (W·F,)  residual stream x", "#FCE4EC"),
            ("Block FC1  in_dim → 64  ReLU + dropout", "#F8BBD0"),
            ("Block FC2  64 → 64  ReLU + dropout", "#F48FB1"),
            ("Block FC3  64 → 64  ReLU + dropout", "#F06292"),
            ("θ_b, θ_f  polynomial heads (degree 8)  × Vandermonde Vb / Vf", "#EC407A"),
            ("Backcast subtract  x ← x − bc    Forecast add  ŷ ← ŷ + fc", "#E91E63"),
            ("Repeat 2 stacks × 3 blocks = 6 residual blocks", "#C2185B"),
        ],
    )
    paths["nbeats"] = os.path.join(DIAG_DIR, "nbeats.png")

    _stack_diagram(
        os.path.join(DIAG_DIR, "transformer.png"),
        "Transformer encoder — layer stack (NumPy, Tab 1)",
        [
            ("Input  (W, F)  multivariate window", "#E0F2F1"),
            ("Linear embed  F → d_model  + sinusoidal PE (length W)", "#B2DFDB"),
            ("MHA-1  n_heads=2  scaled dot-product  residual + dropout", "#80CBC4"),
            ("FFN  d_model→32 ReLU dropout → d_model  residual + dropout", "#4DB6AC"),
            ("MHA-2  second self-attention  residual + dropout", "#26A69A"),
            ("Mean pool over time  (W, d) → (d,)", "#00897B"),
            ("Dense  d_model → horizon  linear  ŷ", "#00695C"),
        ],
    )
    paths["tf"] = os.path.join(DIAG_DIR, "transformer.png")

    _stack_diagram(
        os.path.join(DIAG_DIR, "sarima.png"),
        "SARIMA — statistical architecture (Tab 1 baseline)",
        [
            ("Input  1-D scaled claim series only  (exog not in SARIMAX)", "#F1F8E9"),
            ("Difference  d  + seasonal difference  D  (period 12)", "#DCEDC8"),
            ("AR / SAR lags  p, P     MA / SMA lags  q, Q", "#C5E1A5"),
            ("statespace.SARIMAX  MLE fit  maxiter=100", "#AED581"),
            ("h-step forecast  clip ≥ 0    else seasonal-naive / mean fallback", "#9CCC65"),
        ],
        height=5.0,
    )
    paths["sarima"] = os.path.join(DIAG_DIR, "sarima.png")

    _stack_diagram(
        os.path.join(DIAG_DIR, "keras_cnn_lstm.png"),
        "Keras-CNN-LSTM — Tab 5 annotated path (optional TF)",
        [
            ("Input  (time_step=4, 5 features)  Part_Failure+Production+…", "#EDE7F6"),
            ("Conv1D  64 filters  k=2  ReLU", "#D1C4E9"),
            ("Dropout 0.2", "#B39DDB"),
            ("LSTM 50  return_sequences=True  + Dropout 0.2", "#9575CD"),
            ("LSTM 50  return_sequences=False  + Dropout 0.2", "#7E57C2"),
            ("Flatten  → Dense(1)  MSE + Adam  EarlyStopping / LR schedule", "#673AB7"),
        ],
    )
    paths["keras"] = os.path.join(DIAG_DIR, "keras_cnn_lstm.png")

    _stack_diagram(
        os.path.join(DIAG_DIR, "ensemble.png"),
        "Feature fusion + inverse-MAE ensemble",
        [
            ("Per-t fusion  [claim_t | EXOG_t]  → F channels", "#FFF3E0"),
            ("Each model → 12-mo forecast (scaled)  then inverse MinMax", "#FFE0B2"),
            ("w_k = (1 / CV_MAE_k) / Σ 1/MAE     ŷ_ens = Σ w_k ŷ_k", "#FFCC80"),
            ("Optional CM multiplier (FCO/K share × reduction)", "#FFB74D"),
            ("CI: multi-model std of forecasts  or  0.25×σ(last 12) if one model", "#FFA726"),
        ],
        height=5.2,
    )
    paths["ens"] = os.path.join(DIAG_DIR, "ensemble.png")
    return paths


def _add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def _p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return para


def _code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    para.paragraph_format.space_after = Pt(8)
    return para


def _picture(doc, path, width=6.4):
    if os.path.isfile(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def build_docx(paths: dict[str, str]) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    title = doc.add_heading(
        "Multivariate Forecasting Models — Complete Architecture", 0
    )
    _p(
        doc,
        "Automotive Warranty Claims Forecasting System. This document describes "
        "every multivariate model used in the project: NumPy CNN-LSTM, N-BEATS, "
        "Transformer, SARIMA, the inverse-MAE ensemble, and the optional Keras "
        "CNN-LSTM used on the annotated walk-forward tab. Source of truth: "
        "forecasting/models/*.py, forecasting/pipeline/runner.py, "
        "forecasting/data/loader.py, forecasting/dashboard/app_ui.py.",
    )
    _p(doc, "Generated to match the implementation (not a generic textbook architecture).")

    _add_heading(doc, "1. Scope and model roster", 1)
    _p(
        doc,
        "Multivariate here means the Tab 1 pipeline that forecasts claim counts "
        "from a lookback window that concatenates the target series with exogenous "
        "drivers (production, odometer, vehicle age, FCO/K batch intensity, "
        "Fourier seasonality, calendar, lags). SARIMA is trained on the claim "
        "series only but is scored in the same multivariate pipeline.",
    )
    _table(
        doc,
        ["Model", "Implementation", "Purpose", "Primary code"],
        [
            ["CNN-LSTM", "Pure NumPy", "Local patterns (CNN) + temporal memory (LSTM)", "models/cnn_lstm.py"],
            ["N-BEATS", "Pure NumPy", "Interpretable basis expansion with residual stacks", "models/nbeats.py"],
            ["Transformer", "Pure NumPy", "Long-range dependencies via self-attention", "models/transformer.py"],
            ["SARIMA", "statsmodels", "Classical seasonal baseline", "models/ml_models.py"],
            ["Ensemble", "Pipeline", "Inverse-MAE blend of selected models", "pipeline/runner.py"],
            ["Keras-CNN-LSTM", "TensorFlow optional", "Reference annotated last-6-month test", "pipeline/annotated_forecast.py"],
        ],
    )
    _p(
        doc,
        "Not used in this project: XGBoost, LightGBM, Gradient Boosting. "
        "Tree ML is excluded from MODEL_NAMES in forecasting/config.py.",
    )

    _add_heading(doc, "2. Shared data flow and integration points", 1)
    _picture(doc, paths["pipeline"], 6.6)
    _add_heading(doc, "2.1 Input layer (data, not the neural net)", 2)
    _p(
        doc,
        "Claims rows are validated and normalised (Part Name, FCOK_DATE, "
        "PROCESSING_DATE, ODOMETER). Aggregation is by process month. "
        "Production is user-supplied (synthetic production is disabled for training). "
        "Gaps in the monthly calendar are reindexed and filled.",
    )
    _add_heading(doc, "2.2 Preprocessing and feature extraction (fusion)", 2)
    _p(doc, "EXOG_COLS fused with scaled claims at each time step:")
    _code(
        doc,
        "avg_odometer, avg_vehicle_age, avg_mfctr_age, production,\n"
        "sin_12, cos_12, sin_6, cos_6,\n"
        "median_odometer, max_odometer, median_vehicle_age,\n"
        "n_fcok_batches, warranty_share, Vehicle_Age,\n"
        "Month, Year, Quarter,\n"
        "Claim_Lag_1, Claim_Lag_2, Claim_Lag_3, Claim_Lag_12",
    )
    _p(
        doc,
        "Window builder (runner.build_window_dataset): for each start index i, "
        "X[i] = concat(claim_window, exog_window) along the feature axis; "
        "y[i] = next `horizon` claim values. Shape X: (N, W, 1+F), y: (N, H). "
        "Walk-forward CV uses horizon=1; final refit uses FORECAST_HORIZON=12.",
    )
    _add_heading(doc, "2.3 Normalization", 2)
    _p(
        doc,
        "Two MinMaxScalers: claims and exogenous matrix. Each CV fold fits "
        "scalers only on [:train_end] (no leakage). Inverse transform clips "
        "scaled values to [0, 1] then clips claims to ≥ 0. Full-history refit "
        "fits scalers on all months before the 12-month forecast.",
    )
    _add_heading(doc, "2.4 Shared DL training utilities (base.py)", 2)
    _p(
        doc,
        "Activations: ReLU, numerically stable sigmoid, clipped tanh. "
        "Adam (Kingma & Ba) per-parameter keys. Inverted dropout (scale 1/(1-p) at train). "
        "EarlyStopTracker restores best weights after patience=5 epochs with no loss drop. "
        "RANDOM_SEED=42.",
    )
    _picture(doc, paths["ens"], 6.2)

    _add_heading(doc, "2.5 Ranking, best model, ensemble, CI, countermeasure", 2)
    _p(
        doc,
        "Metrics on walk-forward OOS actuals: RMSE 0.30, MAE 0.25, MAPE 0.25, R² 0.20 "
        "(min-max normalised costs). Best model is auto-selected when multiple models "
        "are checked; a single checked model is used as-is. Ensemble weights are "
        "inverse of mean CV MAE. Config-file CM multiplies the forecast by a decay "
        "curve; Tab 1 optional CM and Tab 4 FCO/K simulator adjust post-model outputs. "
        "CI: if one model, 1.5 × 0.25×std(last 12 claims); if several, 1.5 × std of "
        "model forecasts across months.",
    )
    _add_heading(doc, "2.6 Persistence integration", 2)
    _p(
        doc,
        "outputs/models/*.pkl (scalers + DL objects + weights), "
        "outputs/forecasts/*.csv, outputs/best_params.json (locked HP per part). "
        "CNN-LSTM HP is never grid-searched (fixed lr=1e-2, epochs=40).",
    )

    # CNN-LSTM
    _add_heading(doc, "3. CNN-LSTM", 1)
    _p(doc, "Purpose: extract short local motifs in the multivariate window with 1-D convolutions, then model sequential dependence with an LSTM, and emit a multi-step claim forecast.", bold=False)
    _picture(doc, paths["cnn"], 6.2)
    _add_heading(doc, "3.1 Layer-by-layer", 2)
    _table(
        doc,
        ["Stage", "What happens", "Shapes / params"],
        [
            ["Input", "One sample window of fused features", "(W, F); default W=12"],
            ["Preprocess", "Leak-free MinMax in runner", "F = 1 + n_exog (~21)"],
            ["Conv1", "Valid 1-D conv, kernel 3, 16 filters, ReLU", "(W,F)→(W-2, 16); Wc1 (3,F,16)"],
            ["Dropout", "Inverted dropout rate 0.2 if training", "same as Conv1"],
            ["Conv2", "Valid 1-D conv, kernel 3, 8 filters, ReLU", "(W-2,16)→(W-4, 8)"],
            ["Dropout", "Inverted dropout 0.2", "same as Conv2"],
            ["LSTM", "Scan time; 4 concatenated gates; return last h", "h,c ∈ R^32; Wlstm (8+32, 128)"],
            ["Dropout", "On final hidden state", "R^32"],
            ["Output Dense", "Linear map to horizon", "Wd (32, H); inference H=12"],
        ],
    )
    _p(
        doc,
        "LSTM gates: i=σ, f=σ (bias init 1 on forget), g=tanh candidate, o=σ; "
        "c = f⊙c + i⊙g; h = o⊙tanh(c). Valid convolution shortens the sequence "
        "(12 → 10 → 8 steps into the LSTM).",
    )
    _add_heading(doc, "3.2 Additional features unique to CNN-LSTM", 2)
    bullets = [
        "Forget-gate bias initialised to 1 to reduce vanishing memory at start of training.",
        "Analytic MSE gradients on Dense Wd/bd; sparse numerical gradients (128 random weights) on LSTM tensors only — Conv kernels are not updated (fixed random feature extractor after init).",
        "No HP grid (speed); parallel fold training with other DL models.",
        "On fold exception, pipeline records actual as prediction (degraded fold) — treat as an edge-case signal in ranking.",
        "Dropout off at predict().",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    _add_heading(doc, "3.3 Pseudo-architecture", 2)
    _code(
        doc,
        "x: (W, F)\n"
        "c1 = Dropout(ReLU(Conv1D_k3(x, 16)))\n"
        "c2 = Dropout(ReLU(Conv1D_k3(c1, 8)))\n"
        "h  = LSTM_32(c2)[-1]          # last hidden\n"
        "h  = Dropout(h)\n"
        "y  = Dense(H)(h)              # linear",
    )
    _add_heading(doc, "3.4 Pipeline / UI integration", 2)
    _p(
        doc,
        "Factory _make_dl_model('CNN-LSTM'). CV: horizon 1; refit: horizon 12 on last window. "
        "Appears on Tab 1 ranking, forecast overlay, Tab 2 model comparison / CV folds, "
        "Tab 5 as NumPy CNN-LSTM, Tab 6 ranking CSV / PPT.",
    )

    # N-BEATS
    _add_heading(doc, "4. N-BEATS", 1)
    _p(
        doc,
        "Purpose: forecast via doubly residual stacks of fully connected blocks with "
        "polynomial basis expansion (Oreshkin et al., 2020 style). Multivariate windows "
        "are flattened so every exogenous channel is visible to the first FC layer.",
    )
    _picture(doc, paths["nbeats"], 6.2)
    _add_heading(doc, "4.1 Layer-by-layer (one block)", 2)
    _table(
        doc,
        ["Stage", "What happens", "Shapes / params"],
        [
            ["Input", "Flatten window", "(W·F,); in_dim = lookback * n_features"],
            ["FC1–FC3", "Hidden 64, ReLU, dropout 0.2 each", "W1 (in,64), W2/W3 (64,64)"],
            ["θ heads", "Linear to theta_dim=8 (backcast & forecast)", "Wtb, Wtf (64, 8)"],
            ["Basis", "Vandermonde Vb on [-1,0], Vf on [0,1]", "bc = θ_b V_bᵀ; fc = θ_f V_fᵀ"],
            ["Residual", "x ← x − backcast; ŷ ← ŷ + forecast", "6 blocks (2×3)"],
            ["Output", "Sum of block forecasts", "(H,)"],
        ],
    )
    _add_heading(doc, "4.2 Additional features unique to N-BEATS", 2)
    for b in [
        "Doubly residual topology: backcast cleans the residual stream; forecasts accumulate (feature fusion across blocks).",
        "Polynomial (not Fourier) basis — trend-like expansion on a normalised time grid.",
        "Training updates only forecast heads Wtf/btf via sparse numerical gradients (40 elements/step); other weights stay at init (random features + residual structure).",
        "HP grid: lr ∈ {5e-3, 1e-3}, epochs ∈ {50, 80} on fold-1 windows when tuning is on.",
        "Default constructor epochs=100; runner uses grid / locked params.",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    _add_heading(doc, "4.3 Pseudo-architecture", 2)
    _code(
        doc,
        "x = flatten(window)           # residual\n"
        "y = 0\n"
        "for block in 6_blocks:\n"
        "    h = FC3(Dropout(ReLU(·))) # three layers\n"
        "    bc = (h @ Wtb + btb) @ Vb.T\n"
        "    fc = (h @ Wtf + btf) @ Vf.T\n"
        "    x = x - bc\n"
        "    y = y + fc\n"
        "return y",
    )
    _add_heading(doc, "4.4 Pipeline / UI integration", 2)
    _p(
        doc,
        "Same window fusion as CNN-LSTM. Tab 1 ranking colour #FF6584. "
        "Tab 3 univariate reuses the same class with claims-only UNI_FEATURE_COLS (not this document’s multivariate EXOG set).",
    )

    # Transformer
    _add_heading(doc, "5. Transformer", 1)
    _p(
        doc,
        "Purpose: relate every month in the lookback to every other month through "
        "multi-head self-attention, then pool and map to the horizon. Designed as a "
        "lightweight encoder (no decoder, no causal mask).",
    )
    _picture(doc, paths["tf"], 6.2)
    _add_heading(doc, "5.1 Layer-by-layer", 2)
    _table(
        doc,
        ["Stage", "What happens", "Shapes / params"],
        [
            ["Input", "Multivariate window", "(W, F)"],
            ["Embedding", "Linear F→d_model + PE", "We (F, d); PE (W, d); d default 16"],
            ["Attention", "Per head Q,K,V; softmax(QKᵀ/√dh); concat; Wo", "nh=2, dh=d/2"],
            ["Residual + dropout", "e ← e + Drop(Attn(e))", "(W, d)"],
            ["FFN", "ReLU hidden 32, then d_model, residual + dropout", "Wff1 (d,32), Wff2 (32,d)"],
            ["MHA-2", "Second attention residual (no second FFN)", "(W, d)"],
            ["Pool + Dense", "Mean over time; linear to H", "Wd (d, H)"],
        ],
    )
    _p(
        doc,
        "Softmax is numerically stabilised by subtracting row max. Attention is encoder-style "
        "(full W×W), not causal — the model may attend to all lookback positions equally.",
    )
    _add_heading(doc, "5.2 Additional features unique to Transformer", 2)
    for b in [
        "Sinusoidal positional encoding (10000 base) added after the input projection — the only explicit time stamp inside the net.",
        "Pre-norm is not used; residuals wrap attention and FFN (Post-LN-like add).",
        "HP grid: d_model ∈ {16,32}, lr ∈ {3e-3,1e-3}, epochs ∈ {40,70}; n_heads fixed at 2 in the factory.",
        "Analytic Dense grads use mean of (xWe+be+PE) — an approximation that ignores attention/FFN for Wd.",
        "Sparse numerical grads on Wo, Wff2, We (50 elements). Wq/Wk/Wv are not updated (random attention projections).",
        "d_model must be divisible by n_heads.",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    _add_heading(doc, "5.3 Pseudo-architecture", 2)
    _code(
        doc,
        "e = x @ We + be + PE\n"
        "e = e + Dropout(MHA(e))\n"
        "e = e + Dropout(Wff2(Dropout(ReLU(e @ Wff1))))\n"
        "e = e + Dropout(MHA(e))\n"
        "y = mean(e, axis=time) @ Wd + bd",
    )
    _add_heading(doc, "5.4 Pipeline / UI integration", 2)
    _p(
        doc,
        "PE is built for self.W = lookback at init. Lookback W is shortened when the series is short "
        "(W = min(12, T − horizon − folds − 2)), so the model is reconstructed per part. "
        "Tab 1 / 2 / 5 / 6 same as other DL models.",
    )

    # SARIMA
    _add_heading(doc, "6. SARIMA (multivariate pipeline baseline)", 1)
    _p(
        doc,
        "Purpose: robust seasonal-difference baseline on the claim count series. "
        "It does not consume EXOG inside SARIMAX; integration is at ranking, ensemble, and UI.",
    )
    _picture(doc, paths["sarima"], 6.2)
    _add_heading(doc, "6.1 Layer-by-layer (statistical)", 2)
    _table(
        doc,
        ["Stage", "What happens", "Notes"],
        [
            ["Input", "Scaled 1-D claim_count", "Same MinMax as DL target"],
            ["Preprocess", "d, D differences inside statespace", "enforce_stationarity/invertibility False"],
            ["ARMA", "Non-seasonal (p,d,q)", "Grid: (1,1,1), (1,1,0), (0,1,1)"],
            ["Seasonal", "(P,D,Q,s) s=12", "(1,1,0,12) or (0,1,1,12)"],
            ["Output", "forecast(steps=h), clip ≥ 0", "h=1 in CV; h=12 on refit"],
        ],
    )
    _add_heading(doc, "6.2 Additional features unique to SARIMA", 2)
    for b in [
        "Fallback if MLE fails: seasonal naive last-12 cycle, else series mean.",
        "grid_search_sarima minimises 1-step MAE on fold-1 next month (scaled).",
        "JSON lock converts order tuples via _restore_param_types.",
        "No dropout / attention / residual nets.",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    _add_heading(doc, "6.3 Pseudo-architecture", 2)
    _code(
        doc,
        "try:\n"
        "    m = SARIMAX(y, order=(p,d,q), seasonal_order=(P,D,Q,12))\n"
        "    return clip(m.fit().forecast(H), 0, inf)\n"
        "except:\n"
        "    return seasonal_naive(y) or mean(y)",
    )

    # Ensemble
    _add_heading(doc, "7. Inverse-MAE ensemble", 1)
    _p(
        doc,
        "Purpose: combine selected multivariate models into a single 12-month path "
        "and provide disagreement-based confidence bands. Not a separately trained network.",
    )
    _code(
        doc,
        "w_k = (1 / (MAE_k + 1e-9)) / sum_j (1 / (MAE_j + 1e-9))\n"
        "y_ens_scaled = sum_k w_k * y_k_scaled\n"
        "y_ens = inverse_minmax(y_ens_scaled) * cm_mults",
    )
    _p(
        doc,
        "If the user selects one model, weights collapse to 1.0 on that model and "
        "ensemble_raw equals the best (only) forecast. Integration: result['ensemble_raw'], "
        "forecast table columns, overlay on Tab 1 plot, pickle bundle['weights'].",
    )

    # Keras
    _add_heading(doc, "8. Keras-CNN-LSTM (Tab 5 only)", 1)
    _p(
        doc,
        "Purpose: optional TensorFlow path for the annotated last-6-month evaluation "
        "with a different, smaller feature set. Disabled when TensorFlow is not installed "
        "(typical on Python 3.14). Not used for the Tab 1 12-month production forecast.",
    )
    _picture(doc, paths["keras"], 6.2)
    _table(
        doc,
        ["Stage", "Detail"],
        [
            ["Input features", "Part_Failure, Production, Warranty_Days, Countermeasure, FCOK_Jan_Aug"],
            ["Window", "time_step=4, target = next Part_Failure"],
            ["Conv1D", "64 filters, kernel 2, ReLU"],
            ["LSTM stack", "50 units sequences True, then 50 units last state"],
            ["Regularization", "Dropout 0.2 × 3; EarlyStopping; LR scheduler"],
            ["Output", "Dense(1); inverse scale with zero-padded extra channels"],
        ],
    )

    # UI
    _add_heading(doc, "9. Tab-wise UI outputs (what each tab means)", 1)
    _table(
        doc,
        ["Tab", "Name", "What it shows for multivariate models", "How to read it"],
        [
            ["1", "Upload & Forecast", "Upload, part, model multi-select, production/cost, CM prompts, status (best model), history+forecast plot, 12-mo table, actual vs predicted, ranking, hyperparameters, CPV/claim-ratio history", "Green load = schema OK. Ranking is leak-free OOS. Forecast table is the selected/best model after CM. Production must be positive every month to train."],
            ["2", "Diagnostics", "Production, vehicle age, odometer, FCO/K×process heatmap, rolling CV MAE by fold, model comparison bars", "Heatmap = which manufacturing months drive process-month claims. CV plot = stability. Comparison = RMSE/MAE etc. aligned with ranking."],
            ["3", "Univariate Analysis", "Claims-only Holt-Winters/SARIMA/DL — not the Tab 1 EXOG fusion", "Do not treat Tab 3 metrics as multivariate scores."],
            ["4", "Reduction", "FCO/K month picker, reduction %, original vs adjusted forecast, monthly/cumulative reduction, share, CSV", "What-if on peak batches over 36-month warranty. Slider ~55% sensitivity floor in UI copy. Requires a Tab 1 train first."],
            ["5", "Annotated Walk-Forward", "Last 6 months actual vs forecast, Error, Error%, Accuracy% per model; Keras optional; annotated CSV", "Stress-test feeding. Accuracy% is 1-|e|/actual when actual≠0. Run one model first (slow)."],
            ["6", "Summary & Export", "KPI HTML, summary table, global ranking, forecast/ranking CSV, PPTX", "PPT includes ranking, economics, CM, optional univariate if Tab 3 was run. kaleido for chart images."],
        ],
    )
    _add_heading(doc, "9.1 Tab 1 widgets mapped to models", 2)
    for b in [
        "Models for forecasting checkbox → selected_models → runner._resolve_selected_models.",
        "Tune hyperparameters → HP_TUNE path; else locked outputs/best_params.json.",
        "12-Month Forecast plot → best_forecast + ensemble/CI when multiple models.",
        "Actual vs Predicted → oos_actual vs oos_preds[model] for each fold month.",
        "Hyperparameters Used → best_params including SARIMA tuples and DL lr/epochs/d_model.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    _add_heading(doc, "10. Edge-case testing notes", 1)
    _table(
        doc,
        ["Case", "Expected behaviour", "Where to verify"],
        [
            ["Empty / unknown part", "No monthly rows; train skipped", "Tab 1 status"],
            ["Zero total claims", "run_pipeline_for_part returns None, SKIP", "Log / status"],
            ["Series too short (W<4)", "SKIP — cannot fit lookback + 6 folds + horizon", "Tab 1"],
            ["Missing production / zeros", "ValueError — synthetic production disabled", "Tab 1 table"],
            ["Missing EXOG columns", "use_cols intersection; nan_to_num 0", "result['exog_cols']"],
            ["All-zero OOS actuals", "MAPE NaN (safe_mape); R² 0 if variance 0", "Ranking table"],
            ["SARIMA MLE fail", "Seasonal naive or mean", "Forecast still ≥ 0"],
            ["DL fold crash", "Pred set to actual; MAE 1.0 scaled — ranking distorted", "Tab 2 CV; logs"],
            ["Single model selected", "That model is 'best'; ensemble weight 1; CI from residual scale", "Tab 1 plot band"],
            ["HP tune off + no lock file", "Defaults / empty SARIMA dict → (1,1,1)×(1,1,0,12)", "HP table"],
            ["Transformer d_model=32", "n_heads still 2; dh=16", "HP table"],
            ["Lookback ≠ 12", "CNN valid conv still k=3; Transformer PE resized at init", "Artifacts lookback"],
            ["CM with no FCO/K months (Tab 4)", "Warning: select months", "Tab 4"],
            ["Tab 5 without Tab 1 load", "Status: load data & select part", "Tab 5"],
            ["No TensorFlow", "Keras-CNN-LSTM hidden from checkbox", "Tab 5 choices"],
            ["Inverse scale clip [0,1]", "Scaled forecasts outside [0,1] are clipped — extreme DL outputs saturates", "Compare scaled vs raw"],
            ["Parallel HP / fold threads", "Non-deterministic completion order; seed still set", "Re-runs may differ slightly"],
            ["Tab 3 vs Tab 1", "Different features and often different winners", "Do not mix rankings"],
        ],
    )

    _add_heading(doc, "11. Recommended edge-case test checklist", 1)
    for b in [
        "Happy path: upload template claims + production, all four models, tune on, 12-mo table length = 12, ranking has 4 rows, CI band present.",
        "CNN-LSTM only: ranking 1 row, no multi-model std CI, HP shows lr=0.01 epochs=40.",
        "SARIMA only: no DL pickle keys required; forecast still plots.",
        "Short history (~18 months): either SKIP or reduced W; document the message.",
        "Zero-claim months in the middle: lags filled with mean; heatmap still renders.",
        "Tab 4: one FCO/K month at 0% ≈ original; 100% reduces only that batch’s share, not the entire forecast to zero.",
        "Tab 5: six test months, Accuracy% blank/NaN when actual is 0.",
        "Export: PPT builds after Tab 1; fails clearly if python-pptx missing.",
        "Lock file: second train with tune off reuses JSON; SARIMA orders remain tuples after reload.",
    ]:
        doc.add_paragraph(b, style="List Number")

    _add_heading(doc, "12. File map", 1)
    _table(
        doc,
        ["Path", "Role"],
        [
            ["forecasting/models/base.py", "ReLU/sigmoid/tanh, dropout, early stop, Adam"],
            ["forecasting/models/cnn_lstm.py", "CNN-LSTM"],
            ["forecasting/models/nbeats.py", "N-BEATS blocks"],
            ["forecasting/models/transformer.py", "Transformer encoder"],
            ["forecasting/models/ml_models.py", "SARIMA + HP_GRIDS"],
            ["forecasting/pipeline/runner.py", "Windows, CV, rank, ensemble, persist"],
            ["forecasting/pipeline/annotated_forecast.py", "Tab 5 Keras + NumPy"],
            ["forecasting/data/loader.py", "EXOG, monthly, FCO/K CM"],
            ["forecasting/metrics.py", "RMSE/MAE/MAPE/R² rank"],
            ["forecasting/dashboard/app_ui.py", "Gradio tabs 1–6"],
            ["forecasting/config.py", "MODEL_NAMES, seeds, weights"],
        ],
    )

    os.makedirs(os.path.dirname(DOCX_PATH), exist_ok=True)
    doc.save(DOCX_PATH)
    print("Wrote", DOCX_PATH)


def main():
    paths = generate_diagrams()
    print("Diagrams in", DIAG_DIR)
    try:
        build_docx(paths)
    except ImportError:
        print("python-docx missing — installing is required for .docx")
        raise


if __name__ == "__main__":
    main()
